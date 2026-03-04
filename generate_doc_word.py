#!/usr/bin/env python3
"""
generate_doc_word.py — Carousel Allocation Tool
Documentation Utilisateur — version Word (.docx)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ── Colour palette ────────────────────────────────────────────────────────────
HEX_DARK   = "2C2C2C"   # Hub Performance dark charcoal
HEX_RED    = "C8321A"   # Hub Performance brand red
HEX_LIGHT  = "FEF0EE"   # Light pink tint
HEX_GREY   = "F5F5F5"   # Light grey (alternating rows)
HEX_MID    = "D0D0D0"   # Border grey
HEX_WHITE  = "FFFFFF"
HEX_NOTE   = "555555"

C_DARK   = RGBColor(0x2C, 0x2C, 0x2C)
C_RED    = RGBColor(0xC8, 0x32, 0x1A)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
C_NOTE   = RGBColor(0x55, 0x55, 0x55)

PAGE_W_CM    = 21.0
MARGIN_CM    = 2.5
CONTENT_W_CM = PAGE_W_CM - 2 * MARGIN_CM   # ~16 cm


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Set cell background via XML (no # prefix)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_table_borders(table, color_hex=HEX_MID, size="4"):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_table_borders_none(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_para_shading(para, hex_color):
    """Set paragraph background color."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    pPr.append(shd)


def add_tab_stop_right(para, pos_cm):
    """Add a right-aligned tab stop at pos_cm."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(pos_cm * 567)))
    tabs.append(tab)
    pPr.append(tabs)


def add_bottom_border(para, color_hex=HEX_RED, size="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)


def add_top_border(para, color_hex=HEX_RED, size="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(top)
    pPr.append(pBdr)


# ── HTML-to-runs parser ───────────────────────────────────────────────────────

def _html_segments(html_text):
    """
    Parse basic HTML into list of (text, bold, italic, color|None).
    Handles: <b>, <i>, <font color='#xxx'>; entities &nbsp; &amp;
    """
    text = html_text.replace("&nbsp;", "\u00a0").replace("&amp;", "&")
    bold_stack   = [False]
    italic_stack = [False]
    color_stack  = [None]
    segments     = []
    last         = 0
    tag_re = re.compile(r"<(/?)(\\w+)([^>]*)>", re.IGNORECASE)
    # Use finditer manually to avoid backslash confusion
    tag_re = re.compile(r"<(/?)(\w+)([^>]*)>", re.IGNORECASE)
    for m in tag_re.finditer(text):
        if m.start() > last:
            seg = text[last:m.start()]
            if seg:
                segments.append((seg, bold_stack[-1], italic_stack[-1], color_stack[-1]))
        is_close = m.group(1) == "/"
        tag      = m.group(2).lower()
        attrs    = m.group(3)
        if not is_close:
            if   tag == "b":    bold_stack.append(True)
            elif tag == "i":    italic_stack.append(True)
            elif tag == "font":
                cm = re.search(r"color=['\"]([^'\"]+)['\"]", attrs, re.I)
                if cm:
                    h = cm.group(1).lstrip("#")
                    try:
                        color_stack.append(RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)))
                    except Exception:
                        color_stack.append(color_stack[-1])
                else:
                    color_stack.append(color_stack[-1])
        else:
            if   tag == "b"    and len(bold_stack)   > 1: bold_stack.pop()
            elif tag == "i"    and len(italic_stack)  > 1: italic_stack.pop()
            elif tag == "font" and len(color_stack)   > 1: color_stack.pop()
        last = m.end()
    if last < len(text) and text[last:]:
        segments.append((text[last:], bold_stack[-1], italic_stack[-1], color_stack[-1]))
    return segments


def fill_para(para, html_text, font_size=10, def_bold=False, def_italic=False,
              def_color=None):
    """Add HTML-parsed runs into an existing paragraph object."""
    if def_color is None:
        def_color = C_BLACK
    for seg, bold, italic, color in _html_segments(html_text):
        if seg:
            run = para.add_run(seg)
            run.bold          = bold   or def_bold
            run.italic        = italic or def_italic
            run.font.size     = Pt(font_size)
            run.font.color.rgb = color if color is not None else def_color


# ── Building-block helpers ────────────────────────────────────────────────────

def sp(doc, pt=6):
    """Spacer paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.line_spacing = Pt(pt)
    r = para.add_run()
    r.font.size = Pt(1)
    return para


def add_hr(doc, color=HEX_RED):
    """Thin horizontal line via paragraph bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    add_bottom_border(para, color)
    return para


def add_section_bar(doc, text):
    """Dark charcoal full-width bar with white bold title."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.3)
    set_para_shading(para, HEX_DARK)
    run = para.add_run(text)
    run.bold          = True
    run.font.size     = Pt(12)
    run.font.color.rgb = C_WHITE
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = C_RED
    return para


def add_body(doc, html_text, sba=3, saf=5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sba)
    para.paragraph_format.space_after  = Pt(saf)
    fill_para(para, html_text, font_size=10, def_color=C_BLACK)
    return para


def add_note(doc, html_text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.5)
    fill_para(para, html_text, font_size=9, def_italic=True, def_color=C_NOTE)
    return para


def add_bullet(doc, html_text):
    clean = html_text.lstrip("•\u00a0 ")
    para  = doc.add_paragraph()
    para.paragraph_format.space_before      = Pt(2)
    para.paragraph_format.space_after       = Pt(3)
    para.paragraph_format.left_indent       = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(-0.3)
    fill_para(para, "• " + clean, font_size=10, def_color=C_BLACK)
    return para


def make_table(doc, data, col_widths_cm=None, header_bg=HEX_RED):
    """Styled table — alternating row backgrounds, red header."""
    if not data:
        return None
    rows = len(data)
    cols = len(data[0])
    if col_widths_cm is None:
        col_widths_cm = [CONTENT_W_CM / cols] * cols

    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    set_table_borders(table, HEX_MID, "4")

    for ri in range(rows):
        is_hdr = ri == 0
        bg = header_bg if is_hdr else (HEX_WHITE if ri % 2 == 1 else HEX_GREY)
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            txt = str(data[ri][ci]) if ci < len(data[ri]) else ""
            if is_hdr:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fill_para(para, txt, font_size=9, def_bold=True, def_color=C_WHITE)
            else:
                fill_para(para, txt, font_size=9, def_color=C_BLACK)
    return table


# ── Page layout & header/footer ───────────────────────────────────────────────

def setup_layout(doc):
    section = doc.sections[0]
    section.page_width     = Cm(21.0)
    section.page_height    = Cm(29.7)
    section.left_margin    = Cm(MARGIN_CM)
    section.right_margin   = Cm(MARGIN_CM)
    section.top_margin     = Cm(2.0)
    section.bottom_margin  = Cm(1.8)
    section.different_first_page_header_footer = True

    # ── Cover page: no header / no footer ──
    for p in section.first_page_header.paragraphs:
        p.clear()
    for p in section.first_page_footer.paragraphs:
        p.clear()

    # ── Inner pages header ──
    hdr  = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p.clear()
    para = hdr.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    r1 = para.add_run("Carousel Allocation Tool")
    r1.bold          = True
    r1.font.size     = Pt(9)
    r1.font.color.rgb = C_DARK
    para.add_run("\t")
    r2 = para.add_run("Page ")
    r2.font.size     = Pt(9)
    r2.font.color.rgb = C_NOTE
    # Page number field
    for tag, text in [("begin", None), ("instrText", " PAGE "), ("end", None)]:
        fc = OxmlElement("w:fldChar" if tag != "instrText" else "w:instrText")
        if tag == "instrText":
            fc.set(qn("xml:space"), "preserve")
            fc.text = text
        else:
            fc.set(qn("w:fldCharType"), tag)
        r3 = para.add_run()
        r3.font.size      = Pt(9)
        r3.font.color.rgb = C_NOTE
        r3._r.append(fc)
    add_tab_stop_right(para, CONTENT_W_CM)
    add_bottom_border(para, HEX_RED, "4")

    # ── Inner pages footer ──
    ftr  = section.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p.clear()
    para = ftr.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    rl = para.add_run("Hub Performance")
    rl.bold          = True
    rl.font.size     = Pt(8)
    rl.font.color.rgb = C_DARK
    para.add_run("\t")
    rr = para.add_run("Confidentiel – Usage interne")
    rr.font.size     = Pt(8)
    rr.font.color.rgb = C_NOTE
    add_tab_stop_right(para, CONTENT_W_CM)
    add_top_border(para, HEX_RED, "4")


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Vertical spacer to push content down
    for _ in range(6):
        sp(doc, 8)

    # Logo
    if os.path.exists("assets/logo.png"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(12)
        p.add_run().add_picture("assets/logo.png", width=Cm(8))

    add_hr(doc)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run("Carousel Allocation Tool")
    r.bold          = True
    r.font.size     = Pt(28)
    r.font.color.rgb = C_DARK

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(16)
    r = p.add_run("Documentation Utilisateur")
    r.font.size     = Pt(16)
    r.font.color.rgb = C_RED

    add_hr(doc)

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run("Version 1.0  —  Mars 2026")
    r.italic         = True
    r.font.size      = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


# ── Identity page ─────────────────────────────────────────────────────────────

def add_identity_page(doc):
    sp(doc, 10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run("Fiche d'identification du document")
    r.bold          = True
    r.font.size     = Pt(14)
    r.font.color.rgb = C_DARK

    add_hr(doc)
    sp(doc, 4)

    info = [
        ("Titre",        "Carousel Allocation Tool — Documentation Utilisateur"),
        ("Version",      "1.0"),
        ("Date",         "Mars 2026"),
        ("Auteur",       "Aik Sidi Ahmed"),
        ("Usage",        "Hub Performance"),
        ("Technologie",  "Python (usage interne)"),
        ("Diffusion",    "Confidentiel – Usage interne uniquement"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.autofit = False
    set_table_borders(table, HEX_MID, "4")
    for i, (key, val) in enumerate(info):
        bg = HEX_LIGHT if i % 2 == 0 else HEX_WHITE
        kc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        kc.width = Cm(4)
        vc.width = Cm(12)
        set_cell_shading(kc, bg)
        set_cell_shading(vc, bg)
        kc.paragraphs[0].paragraph_format.space_before = Pt(6)
        kc.paragraphs[0].paragraph_format.space_after  = Pt(6)
        vc.paragraphs[0].paragraph_format.space_before = Pt(6)
        vc.paragraphs[0].paragraph_format.space_after  = Pt(6)
        rk = kc.paragraphs[0].add_run(key)
        rk.bold          = True
        rk.font.size     = Pt(10)
        rk.font.color.rgb = C_DARK
        rv = vc.paragraphs[0].add_run(val)
        rv.font.size     = Pt(10)
        rv.font.color.rgb = C_BLACK

    sp(doc, 10)
    add_hr(doc)
    sp(doc, 4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("Objet du document")
    r.bold          = True
    r.font.size     = Pt(11)
    r.font.color.rgb = C_RED

    add_body(doc,
        "Ce document est le guide de référence de l'outil <b>Carousel Allocation Tool</b>, "
        "destiné aux équipes opérations aéroportuaires. Il décrit l'objectif de l'outil, "
        "les données d'entrée requises, la configuration des carrousels et des règles, "
        "le déroulement pas à pas de l'interface, ainsi que les fichiers de sortie générés.")
    add_body(doc,
        "L'outil est accessible via un navigateur web. Aucune installation n'est requise "
        "de la part de l'utilisateur final.")

    doc.add_page_break()


# ── Table of contents ─────────────────────────────────────────────────────────

def add_toc(doc):
    add_section_bar(doc, "Sommaire")
    sp(doc, 4)

    sections = [
        ("1.",    "Présentation de l'outil",                False),
        ("2.",    "Architecture générale",                   False),
        ("3.",    "Données d'entrée",                        False),
        ("3.1",   "Format de fichier accepté",               True),
        ("3.2",   "Colonnes obligatoires",                   True),
        ("3.3",   "Colonnes optionnelles",                   True),
        ("3.4",   "Valeurs attendues par colonne",           True),
        ("4.",    "Configuration de l'outil",                False),
        ("4.1",   "Paramètres des carrousels",               True),
        ("4.2",   "Paramètres de temps",                     True),
        ("4.3",   "Réajustement",                            True),
        ("4.4",   "Règle 1 – Multi-carrousels",              True),
        ("4.5",   "Règle 2 – Narrow → Wide",                 True),
        ("4.6",   "Règle 3 – Extras (carrousels supp.)",     True),
        ("4.7",   "Ordre de priorité des règles",            True),
        ("4.8",   "Couleurs de planning",                    True),
        ("5.",    "Utilisation pas à pas",                   False),
        ("6.",    "Capacités et contraintes",                False),
        ("7.",    "Résultats générés (Outputs)",             False),
        ("7.1",   "Récapitulatif d'allocation (CSV)",        True),
        ("7.2",   "Timeline Excel",                          True),
        ("7.3",   "Heatmap Excel",                           True),
        ("7.4",   "KPIs affichés à l'écran",                 True),
        ("8.",    "Codes d'erreur et avertissements",        False),
        ("9.",    "Glossaire",                               False),
    ]

    table = doc.add_table(rows=len(sections), cols=1)
    table.autofit = False
    set_table_borders(table, "DDDDDD", "2")

    for i, (num, title, is_sub) in enumerate(sections):
        bg   = HEX_WHITE if i % 2 == 0 else HEX_GREY
        cell = table.rows[i].cells[0]
        cell.width = Cm(CONTENT_W_CM)
        set_cell_shading(cell, bg)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        if is_sub:
            para.paragraph_format.left_indent = Cm(0.8)
            r = para.add_run(f"{num}  {title}")
            r.font.size     = Pt(9)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            r = para.add_run(f"{num}  {title}")
            r.bold          = True
            r.font.size     = Pt(10)
            r.font.color.rgb = C_DARK

    sp(doc, 4)
    doc.add_page_break()


# ── Section 1 & 2 ─────────────────────────────────────────────────────────────

def section_presentation(doc):
    add_section_bar(doc, "1.  Présentation de l'outil")
    sp(doc)
    add_body(doc,
        "Le <b>Carousel Allocation Tool</b> est un outil web dédié aux opérations aéroportuaires. "
        "Il automatise l'affectation des vols aux carrousels bagages d'un aéroport en tenant compte "
        "des contraintes de capacité, des catégories d'appareils et des fenêtres horaires de traitement bagages.")
    add_body(doc, "L'outil permet de :")
    add_bullet(doc, "Importer un fichier de vols (Excel ou CSV).")
    add_bullet(doc, "Associer automatiquement chaque vol à un ou plusieurs carrousels disponibles.")
    add_bullet(doc, "Appliquer des règles métier configurables (round-robin, multi-carrousel, réajustement…).")
    add_bullet(doc, "Générer des rapports de résultats : récapitulatif CSV, timeline et heatmap Excel, KPIs.")
    sp(doc)
    add_note(doc,
        "L'outil fonctionne grâce à une application Python déployée en interne — "
        "aucune installation n'est requise de la part de l'utilisateur final.")
    sp(doc, 10)

    add_section_bar(doc, "2.  Architecture générale")
    sp(doc)
    add_body(doc,
        "L'outil se compose d'une interface web accessible via navigateur et d'un moteur de calcul "
        "côté serveur. L'utilisateur interagit uniquement avec l'interface — le traitement est réalisé "
        "en arrière-plan.")
    sp(doc)
    make_table(doc, [
        ["Couche",              "Rôle"],
        ["Interface web",       "Formulaire de saisie, affichage des résultats et téléchargement des fichiers"],
        ["Serveur API",         "Réception des fichiers, gestion des sessions, orchestration des calculs"],
        ["Moteur d'allocation", "Application des règles, calcul des affectations, génération des fichiers de sortie"],
    ], col_widths_cm=[4.5, 11.5])
    sp(doc, 8)
    doc.add_page_break()


# ── Section 3 ─────────────────────────────────────────────────────────────────

def section_input(doc):
    add_section_bar(doc, "3.  Données d'entrée")
    sp(doc)

    add_h2(doc, "3.1  Format de fichier accepté")
    make_table(doc, [
        ["Format", "Extension", "Remarques"],
        ["Excel",  ".xlsx",     "Format recommandé — feuille active lue par défaut"],
        ["CSV",    ".csv",      "Séparateur virgule ou point-virgule détecté automatiquement"],
    ], col_widths_cm=[3.5, 3.0, 9.5])
    sp(doc)
    add_body(doc,
        "Le fichier doit contenir <b>une ligne d'en-tête</b> (noms de colonnes) en première ligne. "
        "Les colonnes peuvent être dans n'importe quel ordre — l'outil propose une étape de mapping "
        "pour les associer aux champs attendus.")

    sp(doc, 8)
    add_h2(doc, "3.2  Colonnes obligatoires")
    add_body(doc,
        "Les quatre colonnes suivantes sont <b>indispensables</b>. Sans elles, l'allocation ne peut pas démarrer.")
    make_table(doc, [
        ["Nom interne",   "Description",                                        "Exemple"],
        ["DepartureTime", "Date et heure de départ du vol",                     "2024-06-15 08:30"],
        ["FlightNumber",  "Identifiant unique du vol",                          "AF1234"],
        ["Category",      "Catégorie de l'appareil : Wide body ou Narrow body", "Wide"],
        ["Positions",     "Nombre de positions bagages requises",               "3"],
    ], col_widths_cm=[3.8, 8.0, 4.2])

    sp(doc, 8)
    add_h2(doc, "3.3  Colonnes optionnelles")
    add_body(doc,
        "Ces colonnes enrichissent l'allocation mais ne sont pas bloquantes si elles sont absentes.")
    make_table(doc, [
        ["Nom interne",   "Description",                                          "Comportement si absente"],
        ["Terminal",      "Terminal de l'aéroport (ex : T1, T2)",                 "Aucun filtrage par terminal"],
        ["MakeupOpening", "Heure d'ouverture du tapis bagages",                   "Calculée via offset de temps"],
        ["MakeupClosing", "Heure de fermeture du tapis bagages",                  "Calculée via offset de temps"],
    ], col_widths_cm=[3.8, 6.5, 5.7])

    sp(doc, 8)
    add_h2(doc, "3.4  Valeurs attendues par colonne")
    make_table(doc, [
        ["Colonne",       "Type",     "Format / Valeurs acceptées",                        "Remarques"],
        ["DepartureTime", "Datetime", "YYYY-MM-DD HH:MM  ou  DD/MM/YYYY HH:MM",            "Doit être parsable comme date"],
        ["FlightNumber",  "Texte",    "Chaîne libre (ex : AF1234, EK 505)",                 "Doit être unique par plage horaire"],
        ["Category",      "Texte",    "\"Wide\" / \"Narrow\" (insensible à la casse)",      "Tout autre valeur génère une alerte"],
        ["Positions",     "Entier",   "Nombre entier positif >= 1",                         "Doit être <= capacité max du carrousel"],
        ["Terminal",      "Texte",    "Chaîne libre correspondant au terminal configuré",   "Ignorée si non configurée"],
        ["MakeupOpening", "Datetime", "Même format que DepartureTime",                      "Optionnel — peut être calculé"],
        ["MakeupClosing", "Datetime", "Même format que DepartureTime",                      "Optionnel — peut être calculé"],
    ], col_widths_cm=[3.2, 2.0, 5.8, 5.0])
    sp(doc)
    add_note(doc,
        "Toute colonne supplémentaire présente dans le fichier d'entrée est conservée et retransmise "
        "dans les fichiers de sortie sans modification.")

    doc.add_page_break()


# ── Section 4 ─────────────────────────────────────────────────────────────────

def section_config(doc):
    add_section_bar(doc, "4.  Configuration de l'outil")
    sp(doc)

    add_h2(doc, "4.1  Paramètres des carrousels")
    add_body(doc,
        "Pour chaque carrousel, l'utilisateur définit sa capacité en positions selon la catégorie d'appareil :")
    make_table(doc, [
        ["Paramètre",        "Description",                                                    "Exemple"],
        ["Identifiant",      "Nom ou numéro du carrousel (ex : C1, C2, Belt 3…)",              "C1"],
        ["Wide positions",   "Nombre de positions disponibles pour les appareils Wide body",   "4"],
        ["Narrow positions", "Nombre de positions disponibles pour les appareils Narrow body", "6"],
    ], col_widths_cm=[4.0, 9.0, 3.0])
    sp(doc)
    add_body(doc,
        "Un appareil <b>Wide body</b> peut utiliser des positions Wide <i>et</i> des positions Narrow "
        "(si configuré). Un appareil <b>Narrow body</b> n'utilise que des positions Narrow.")

    sp(doc, 8)
    add_h2(doc, "4.2  Paramètres de temps")
    make_table(doc, [
        ["Paramètre",             "Description",                                                       "Défaut"],
        ["Time step",             "Granularité de la timeline (en minutes)",                           "5 min"],
        ["Makeup opening offset", "Décalage avant le départ pour l'ouverture du tapis (en minutes)",  "–90 min"],
        ["Makeup closing offset", "Décalage avant le départ pour la fermeture du tapis (en minutes)", "–30 min"],
    ], col_widths_cm=[4.8, 8.2, 3.0])
    sp(doc)
    add_note(doc,
        "Exemple : pour un vol à 10h00 avec offset ouverture = –90 min et fermeture = –30 min, "
        "le carrousel sera occupé de 08h30 à 09h30.")

    doc.add_page_break()

    # ── 4.3 – 4.6 ──
    add_section_bar(doc, "4.  Configuration de l'outil (suite) — Réajustement & Règles")
    sp(doc)

    add_h2(doc, "4.3  Réajustement")
    add_body(doc,
        "Le <b>réajustement</b> est une phase d'optimisation qui s'exécute <i>après</i> l'allocation initiale "
        "(round-robin). Il prend en charge les vols non assignés lors du premier passage et tente de leur "
        "trouver un carrousel en appliquant successivement les règles activées.")
    add_body(doc,
        "Lorsque la case <b>« Appliquer les règles de réajustement »</b> est décochée, "
        "aucune règle complémentaire n'est appliquée — seul le round-robin de base est utilisé. "
        "Les vols qui ne trouvent pas de place restent non assignés avec la raison NO_CAPACITY.")
    add_body(doc,
        "Chaque règle est indépendante. Pour chaque vol non assigné, l'outil essaie les règles "
        "dans l'ordre de priorité défini et s'arrête dès qu'une règle parvient à affecter le vol.")

    sp(doc, 8)
    add_h2(doc, "4.4  Règle 1 – Multi-carrousels")
    add_body(doc,
        "Par défaut, chaque vol est affecté à <b>un seul carrousel</b>. Si la demande en positions "
        "dépasse la capacité d'un carrousel seul, le vol reste non assigné. La règle Multi-carrousels "
        "autorise un vol à être <b>réparti sur plusieurs carrousels simultanément</b> pour absorber "
        "cette demande.")
    make_table(doc, [
        ["Paramètre",                      "Description",                                                          "Valeur par défaut"],
        ["MAX_CAROUSELS_PER_FLIGHT_NARROW", "Nombre maximum de carrousels utilisables pour un vol Narrow body",    "3"],
        ["MAX_CAROUSELS_PER_FLIGHT_WIDE",   "Nombre maximum de carrousels utilisables pour un vol Wide body",      "2"],
    ], col_widths_cm=[6.0, 7.0, 3.0])
    sp(doc)
    add_note(doc,
        "Exemple : un vol Narrow body nécessitant 7 positions peut être réparti sur 3 carrousels "
        "(ex : 3 + 3 + 1 positions). Le résultat est visible dans la colonne <b>CarouselList</b> "
        "du fichier CSV de sortie et affiché avec la couleur Split dans le planning.")

    sp(doc, 8)
    add_h2(doc, "4.5  Règle 2 – Narrow → Wide")
    add_body(doc,
        "Lorsqu'un vol Narrow body ne peut pas être placé sur les positions Narrow disponibles, "
        "cette règle autorise l'outil à le traiter comme un vol <b>Wide body</b> pour l'allocation. "
        "Le vol peut alors occuper des positions Wide body libres.")
    make_table(doc, [
        ["Comportement",            "Description"],
        ["Détection",               "S'applique uniquement si aucune solution Narrow body n'a été trouvée"],
        ["Conversion",              "La catégorie du vol est temporairement changée en Wide pour trouver un carrousel"],
        ["Colonne CategoryChanged", "Marquée « YES » dans le CSV de sortie si la conversion a eu lieu"],
        ["Colonne FinalCategory",   "Indique « Wide » (catégorie après conversion)"],
        ["Couleur planning",        "Affiché en vert (#00B894) dans la timeline et la heatmap — couleur prioritaire"],
    ], col_widths_cm=[4.5, 11.5])
    sp(doc)
    add_note(doc,
        "Cette règle est désactivée par défaut. Elle est utile en cas de saturation des tapis "
        "Narrow body avec des tapis Wide body encore disponibles.")

    sp(doc, 8)
    add_h2(doc, "4.6  Règle 3 – Extras (Carrousels supplémentaires)")
    add_body(doc,
        "Lorsqu'un vol ne peut être placé sur aucun des carrousels existants, la règle Extras crée "
        "dynamiquement des <b>carrousels supplémentaires temporaires</b> (nommés EXTRA1, EXTRA2, etc.) "
        "avec une capacité standard définie par l'utilisateur.")
    add_body(doc,
        "Ces carrousels extras ne font pas partie de la configuration initiale — ils représentent "
        "une capacité d'urgence ou de débordement. Ils apparaissent dans tous les fichiers de sortie "
        "comme des carrousels normaux.")
    make_table(doc, [
        ["Paramètre",         "Description",                                                                      "Valeur par défaut"],
        ["E-Wide capacity",   "Nombre de positions Wide de chaque carrousel extra",                               "Max Wide des carrousels existants (défaut : 8)"],
        ["E-Narrow capacity", "Nombre de positions Narrow de chaque carrousel extra",                             "Max Narrow des carrousels existants (défaut : 4)"],
        ["Nommage",           "Les extras sont créés à la demande : EXTRA1, EXTRA2…",                             "Automatique"],
        ["Par terminal",      "Si des terminaux sont configurés, la capacité extra peut être définie par terminal","Oui, configurable"],
    ], col_widths_cm=[4.0, 7.5, 4.5])
    sp(doc)
    add_note(doc,
        "La capacité par défaut des extras est calculée automatiquement comme le maximum des capacités "
        "des carrousels existants. Elle peut être ajustée manuellement dans l'interface avant de lancer "
        "l'allocation.")

    doc.add_page_break()

    # ── 4.7 – 4.8 ──
    add_section_bar(doc, "4.  Configuration de l'outil (suite) — Priorité & Couleurs")
    sp(doc)

    add_h2(doc, "4.7  Ordre de priorité des règles")
    add_body(doc,
        "Lorsque plusieurs règles sont activées, l'interface propose trois menus déroulants "
        "<b>Priorité 1 / Priorité 2 / Priorité 3</b> qui définissent l'ordre dans lequel les règles "
        "sont tentées pour chaque vol non assigné.")
    add_body(doc,
        "Pour chaque vol non assigné après le round-robin, l'outil tente la règle de Priorité 1 d'abord. "
        "Si elle réussit, les priorités 2 et 3 ne sont pas essayées. "
        "Si elle échoue, il passe à la Priorité 2, puis à la Priorité 3.")
    make_table(doc, [
        ["Priorité",   "Règle recommandée",  "Raisonnement"],
        ["Priorité 1", "Multi-carrousels",   "Essaie d'abord de placer le vol sur les carrousels existants en le répartissant"],
        ["Priorité 2", "Narrow → Wide",      "Si aucun carrousel Narrow n'est disponible, tente les carrousels Wide"],
        ["Priorité 3", "Extras",             "En dernier recours, crée un carrousel supplémentaire"],
    ], col_widths_cm=[2.5, 4.0, 9.5])
    sp(doc)
    add_note(doc,
        "L'ordre peut être modifié selon les contraintes opérationnelles. "
        "Seules les règles activées apparaissent dans les menus de priorité.")

    sp(doc, 8)
    add_h2(doc, "4.8  Couleurs de planning")
    add_body(doc,
        "L'outil affecte une couleur à chaque vol dans la timeline et la heatmap selon son type "
        "d'affectation. Ces couleurs sont personnalisables dans l'interface.")
    make_table(doc, [
        ["Couleur",       "Code hex", "Signification",                                          "Priorité d'affichage"],
        ["Wide",          "#D32F2F",  "Vol Wide body affecté normalement",                      "Normale"],
        ["Narrow",        "#FFCDD2",  "Vol Narrow body affecté normalement",                    "Normale"],
        ["Split",         "#FFC107",  "Vol réparti sur plusieurs carrousels (multi-carrousel)",  "Prioritaire — écrase Wide/Narrow"],
        ["Narrow → Wide", "#00B894",  "Vol Narrow converti et placé sur un carrousel Wide",     "Prioritaire — écrase Wide/Narrow"],
    ], col_widths_cm=[2.5, 2.5, 7.5, 3.5])
    sp(doc)
    add_note(doc,
        "Les couleurs Split et Narrow → Wide ont la priorité sur les couleurs de base Wide et Narrow. "
        "Un vol split sera toujours affiché en jaune, même s'il est de catégorie Wide.")

    doc.add_page_break()


# ── Section 5 ─────────────────────────────────────────────────────────────────

def section_steps(doc):
    add_section_bar(doc, "5.  Utilisation pas à pas")
    sp(doc)
    add_body(doc,
        "L'interface guide l'utilisateur à travers un assistant en 5 étapes. "
        "Chaque étape doit être validée avant de passer à la suivante.")
    sp(doc)

    steps = [
        ("Étape 1 — Import du fichier", [
            "Cliquer sur <b>« Choisir un fichier »</b> et sélectionner le fichier Excel (.xlsx) ou CSV.",
            "L'outil lit automatiquement l'en-tête et affiche un aperçu des premières lignes.",
            "Vérifier que le fichier s'affiche correctement avant de continuer.",
        ]),
        ("Étape 2 — Mapping des colonnes", [
            "L'outil affiche la liste des colonnes détectées dans le fichier.",
            "Pour chaque champ obligatoire (<b>DepartureTime, FlightNumber, Category, Positions</b>), "
            "sélectionner la colonne correspondante dans le fichier.",
            "Les colonnes optionnelles (Terminal, MakeupOpening, MakeupClosing) peuvent être mappées ou ignorées.",
            "Un message d'erreur apparaît si une colonne obligatoire n'est pas mappée.",
        ]),
        ("Étape 3 — Configuration des carrousels", [
            "Renseigner pour chaque carrousel son identifiant, ses positions Wide et ses positions Narrow.",
            "Ajouter autant de carrousels que nécessaire via le bouton <b>« Ajouter un carrousel »</b>.",
            "Si des terminaux sont présents dans le fichier, associer chaque terminal à un groupe de carrousels.",
        ]),
        ("Étape 4 — Paramètres & Règles", [
            "Définir le <b>time step</b> (granularité de la timeline, en minutes).",
            "Renseigner les offsets de temps pour les fenêtres makeup (ouverture / fermeture).",
            "Activer ou désactiver les règles de gestion selon les besoins opérationnels.",
            "Valider les paramètres pour lancer l'allocation.",
        ]),
        ("Étape 5 — Résultats & Export", [
            "L'outil affiche les KPIs : nombre de vols traités, assignés, non assignés, taux d'utilisation.",
            "Un tableau récapitulatif liste chaque vol avec le carrousel affecté et les éventuelles alertes.",
            "Télécharger les fichiers générés : <b>Summary CSV</b>, <b>Timeline Excel</b>, <b>Heatmap Excel</b>.",
            "En cas de vols non assignés, consulter la colonne <b>Raison</b> pour identifier la cause.",
        ]),
    ]

    for num, (title, bullets) in enumerate(steps, 1):
        # Step header table (number + title in red)
        t = doc.add_table(rows=1, cols=2)
        t.autofit = False
        set_table_borders_none(t)
        nc = t.rows[0].cells[0]
        tc = t.rows[0].cells[1]
        nc.width = Cm(1.2)
        tc.width = Cm(CONTENT_W_CM - 1.2)
        set_cell_shading(nc, HEX_RED)
        set_cell_shading(tc, HEX_RED)
        set_cell_borders_none(nc)
        set_cell_borders_none(tc)
        nc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pn = nc.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pn.paragraph_format.space_before = Pt(5)
        pn.paragraph_format.space_after  = Pt(5)
        rn = pn.add_run(str(num))
        rn.bold = True; rn.font.size = Pt(14); rn.font.color.rgb = C_WHITE
        pt = tc.paragraphs[0]
        pt.paragraph_format.space_before = Pt(5)
        pt.paragraph_format.space_after  = Pt(5)
        rt = pt.add_run(title)
        rt.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = C_WHITE

        # Bullet points with light background
        for b in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.space_before      = Pt(2)
            p.paragraph_format.space_after       = Pt(3)
            p.paragraph_format.left_indent       = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            set_para_shading(p, HEX_LIGHT)
            fill_para(p, "• " + b, font_size=10, def_color=C_BLACK)

        sp(doc, 4)

    doc.add_page_break()


# ── Section 6 ─────────────────────────────────────────────────────────────────

def section_capacity(doc):
    add_section_bar(doc, "6.  Capacités et contraintes")
    sp(doc)

    add_h2(doc, "Limites de l'outil")
    make_table(doc, [
        ["Paramètre",                                         "Valeur / Comportement"],
        ["Nombre de vols",                                    "Pas de limite stricte — performances optimales jusqu'à plusieurs milliers de vols par fichier"],
        ["Nombre de carrousels",                              "Pas de limite — autant que nécessaire"],
        ["Positions max par carrousel",                       "Défini par l'utilisateur — pas de plafond technique"],
        ["Vols nécessitant + de positions que la capacité max", "Refusés avec le code NO_CAPACITY ou IMPOSSIBLE_DEMAND"],
        ["Vols hors fenêtre horaire valide",                  "Refusés avec le code BAD_TIME"],
        ["Taille du fichier d'entrée",                        "Pas de limite fixe — les opérations sont vectorisées"],
    ], col_widths_cm=[7.0, 9.0])

    sp(doc, 8)
    add_h2(doc, "Contraintes de capacité")
    add_body(doc, "La capacité d'un carrousel est modélisée avec deux compteurs indépendants :")
    add_bullet(doc, "<b>Wide positions</b> : positions réservées aux appareils Wide body.")
    add_bullet(doc, "<b>Narrow positions</b> : positions réservées aux appareils Narrow body.")
    sp(doc)
    add_body(doc,
        "À chaque instant (selon le time step), la somme des positions occupées ne peut pas dépasser "
        "la capacité déclarée du carrousel. Si la règle <b>Narrow-to-Wide</b> est activée, "
        "un vol Narrow peut consommer des positions Wide si les positions Narrow sont saturées.")

    sp(doc, 8)
    add_h2(doc, "Contraintes temporelles")
    add_body(doc,
        "Chaque vol occupe un carrousel pendant toute la durée de sa fenêtre makeup "
        "(de MakeupOpening à MakeupClosing). Deux vols ne peuvent pas occuper les mêmes positions "
        "sur le même carrousel pendant une même tranche horaire.")

    doc.add_page_break()


# ── Section 7 ─────────────────────────────────────────────────────────────────

def section_outputs(doc):
    add_section_bar(doc, "7.  Résultats générés (Outputs)")
    sp(doc)

    add_h2(doc, "7.1  Récapitulatif d'allocation (CSV)")
    add_body(doc,
        "Fichier CSV téléchargeable listant tous les vols du fichier d'entrée avec les informations d'affectation.")
    make_table(doc, [
        ["Colonne ajoutée",  "Description"],
        ["Carousel",         "Identifiant du carrousel affecté (ex : C1, C2…)"],
        ["CarouselList",     "Liste des carrousels si affectation multi-carrousel (ex : C1;C3)"],
        ["AllocationStatus", "ASSIGNED / UNASSIGNED"],
        ["UnassignedReason", "Code de raison si non assigné (voir section 8)"],
        ["RuleApplied",      "Règle ayant permis l'affectation (round_robin, multi, extra…)"],
    ], col_widths_cm=[5.0, 11.0])
    sp(doc)
    add_note(doc, "Toutes les colonnes du fichier d'entrée sont conservées.")

    sp(doc, 8)
    add_h2(doc, "7.2  Timeline Excel")
    add_body(doc,
        "Fichier Excel représentant l'occupation de chaque carrousel sur l'ensemble de la journée, "
        "découpée selon le time step configuré.")
    make_table(doc, [
        ["Contenu",              "Description"],
        ["Lignes",               "Un carrousel par ligne"],
        ["Colonnes",             "Une tranche horaire par colonne (selon le time step)"],
        ["Valeur de cellule",    "Nombre de positions occupées à cet instant"],
        ["Couleur de cellule",   "Code couleur indiquant le taux d'occupation (vert → orange → rouge)"],
        ["Onglet supplémentaire","Capacité libre restante par carrousel et par tranche horaire"],
    ], col_widths_cm=[4.5, 11.5])

    sp(doc, 8)
    add_h2(doc, "7.3  Heatmap Excel")
    add_body(doc,
        "Fichier Excel de visualisation graphique de l'occupation globale de tous les carrousels. "
        "Permet d'identifier rapidement les pics de charge et les carrousels sous-utilisés.")
    make_table(doc, [
        ["Contenu",            "Description"],
        ["Heatmap globale",    "Vue d'ensemble de l'occupation sur la journée"],
        ["Graphiques",         "Courbes d'occupation par carrousel"],
        ["Taux d'utilisation", "Pourcentage d'utilisation moyen par carrousel sur la journée"],
    ], col_widths_cm=[4.5, 11.5])

    sp(doc, 8)
    add_h2(doc, "7.4  KPIs affichés à l'écran")
    make_table(doc, [
        ["KPI",                       "Description"],
        ["Total vols",                "Nombre total de vols dans le fichier d'entrée"],
        ["Vols assignés",             "Nombre de vols affectés à un carrousel avec succès"],
        ["Vols non assignés",         "Nombre de vols non affectés (avec raison)"],
        ["Taux d'assignation",        "Pourcentage de vols assignés sur le total"],
        ["Taux d'utilisation moyen",  "Occupation moyenne des carrousels sur la période"],
        ["Alertes & avertissements",  "Nombre de problèmes détectés dans les données ou l'allocation"],
    ], col_widths_cm=[5.5, 10.5])

    doc.add_page_break()


# ── Section 8 ─────────────────────────────────────────────────────────────────

def section_errors(doc):
    add_section_bar(doc, "8.  Codes d'erreur et avertissements")
    sp(doc)
    add_body(doc,
        "Lorsqu'un vol ne peut pas être assigné, un code de raison est renseigné dans la colonne "
        "<b>UnassignedReason</b> du fichier CSV de sortie. Ces codes permettent de comprendre "
        "rapidement la cause du problème.")
    sp(doc)
    make_table(doc, [
        ["Code",              "Signification",                                                                   "Action recommandée"],
        ["NO_CAPACITY",       "Aucun carrousel n'avait de positions disponibles pendant la fenêtre du vol",      "Ajouter un carrousel ou activer la règle Extras"],
        ["IMPOSSIBLE_DEMAND", "Le vol demande plus de positions que la capacité maximale de n'importe quel carrousel", "Vérifier la colonne Positions ou augmenter la capacité des carrousels"],
        ["BAD_TIME",          "La fenêtre horaire du vol est invalide (heure de fermeture <= heure d'ouverture)", "Vérifier les colonnes MakeupOpening et MakeupClosing"],
        ["MISSING_COLUMN",    "Une colonne obligatoire est absente ou non mappée",                               "Vérifier l'étape de mapping des colonnes"],
        ["INVALID_CATEGORY",  "La valeur de la colonne Category n'est pas reconnue",                            "Vérifier que les valeurs sont \"Wide\" ou \"Narrow\""],
    ], col_widths_cm=[3.5, 7.0, 5.5])

    sp(doc, 8)
    add_h2(doc, "Avertissements non bloquants")
    add_body(doc,
        "Certaines anomalies génèrent un avertissement visible dans les KPIs mais ne bloquent pas "
        "l'allocation :")
    add_bullet(doc, "Doublons de numéro de vol sur une même plage horaire.")
    add_bullet(doc, "Colonnes optionnelles absentes (Terminal, MakeupOpening, MakeupClosing).")
    add_bullet(doc, "Valeurs manquantes (NaN) dans des colonnes non obligatoires.")

    doc.add_page_break()


# ── Section 9 ─────────────────────────────────────────────────────────────────

def section_glossary(doc):
    add_section_bar(doc, "9.  Glossaire")
    sp(doc)
    make_table(doc, [
        ["Terme",                "Définition"],
        ["Carrousel",            "Tapis roulant bagages sur lequel les passagers récupèrent leurs bagages à l'arrivée ou sur lequel les bagages sont déposés au départ"],
        ["Wide body",            "Catégorie d'avion gros porteur (ex : Boeing 777, Airbus A330, A380) — généralement un grand volume de bagages"],
        ["Narrow body",          "Catégorie d'avion monocouloir (ex : Boeing 737, Airbus A320) — volume de bagages moindre"],
        ["Makeup / Tapis makeup","Période pendant laquelle les bagages d'un vol sont déposés sur le carrousel avant le départ"],
        ["MakeupOpening",        "Heure à laquelle le carrousel est ouvert pour commencer à recevoir les bagages du vol"],
        ["MakeupClosing",        "Heure à laquelle le carrousel est fermé (fin de dépôt des bagages)"],
        ["Positions",            "Nombre d'emplacements occupés par un vol sur un carrousel (lié à la taille de l'avion)"],
        ["Round-robin",          "Algorithme de distribution équitable qui alterne entre les carrousels disponibles"],
        ["Time step",            "Granularité temporelle utilisée pour découper la journée et vérifier les chevauchements"],
        ["KPI",                  "Key Performance Indicator — indicateur clé de performance (ex : taux d'assignation)"],
        ["Multi-carrousel",      "Affectation d'un seul vol sur plusieurs carrousels simultanément si ses besoins dépassent la capacité d'un seul carrousel"],
        ["Heatmap",              "Représentation visuelle par code couleur de l'intensité d'occupation des carrousels"],
        ["Timeline",             "Vue chronologique de l'occupation de chaque carrousel sur la journée"],
    ], col_widths_cm=[4.0, 12.0])

    sp(doc, 10)
    add_hr(doc)
    sp(doc, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("Document rédigé par Aik Sidi Ahmed — Carousel Allocation Tool v1.0 — Mars 2026")
    r.italic          = True
    r.font.size       = Pt(8)
    r.font.color.rgb  = C_NOTE


# ── Build ──────────────────────────────────────────────────────────────────────

def build():
    output_path = "Carousel_Allocation_Tool_Documentation.docx"
    doc = Document()

    # Remove default empty paragraph style noise
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    setup_layout(doc)
    add_cover_page(doc)
    add_identity_page(doc)
    add_toc(doc)
    section_presentation(doc)
    section_input(doc)
    section_config(doc)
    section_steps(doc)
    section_capacity(doc)
    section_outputs(doc)
    section_errors(doc)
    section_glossary(doc)

    doc.save(output_path)
    print(f"OK  Word genere : {output_path}")


if __name__ == "__main__":
    build()
